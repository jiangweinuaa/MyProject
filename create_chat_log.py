#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建聊天记录 Word 文档 - JW_20260311
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

# 创建文档
doc = Document()

# 设置样式
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)

# 标题
title = doc.add_heading('OpenClaw 龙虾助手 - 聊天记录', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 基本信息
doc.add_paragraph()
info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info_para.add_run(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
info_run.font.size = Pt(10)
info_run = info_para.add_run(f'日期：2026-03-11\n')
info_run.font.size = Pt(10)
info_run = info_para.add_run(f'用户：jiangweinuaa\n')
info_run.font.size = Pt(10)
info_run = info_para.add_run(f'助手：龙虾 🦞\n')
info_run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph()

# 目录
doc.add_heading('目录', level=1)
doc.add_paragraph('1. 身份确认与环境检查', style='Heading 2')
doc.add_paragraph('2. GitHub 配置', style='Heading 2')
doc.add_paragraph('3. Java 开发环境安装', style='Heading 2')
doc.add_paragraph('4. MySQL 数据库安装与配置', style='Heading 2')
doc.add_paragraph('5. CRM 微服务开发', style='Heading 2')
doc.add_paragraph('6. Vue 前端开发', style='Heading 2')

doc.add_page_break()

# 聊天记录
chat_records = [
    {
        'time': '08:43',
        'user': '用户',
        'content': '你是 OpenClaw 么？',
        'section': '一、身份确认与环境检查'
    },
    {
        'time': '08:43',
        'user': '龙虾',
        'content': '是的，我是 OpenClaw —— 一个运行在你机器上的 AI 助手。🤖\n\n我刚刚启动，还在熟悉这个环境。你是第一个和我对话的人，所以我们可以一起定义我是谁、怎么称呼你、以及我希望如何帮助你。',
        'section': '一、身份确认与环境检查'
    },
    {
        'time': '08:43',
        'user': '用户',
        'content': '服务器上 Node.js 安装了没，是什么版本的？',
        'section': '一、身份确认与环境检查'
    },
    {
        'time': '08:43',
        'user': '龙虾',
        'content': '安装了，Node.js 版本是 v24.14.0。\n\nnpm 版本：11.9.0\n\n都是很新的版本，环境看起来没问题。',
        'section': '一、身份确认与环境检查'
    },
    {
        'time': '08:44',
        'user': '用户',
        'content': 'Github 账号配置了没',
        'section': '二、GitHub 配置'
    },
    {
        'time': '08:46',
        'user': '龙虾',
        'content': 'GitHub 账号还没有配置。\n\n检查结果：\n- ✅ Git 已安装：v2.43.7\n- ❌ 未配置用户信息\n- ❌ 没有 SSH key',
        'section': '二、GitHub 配置'
    },
    {
        'time': '08:46',
        'user': '用户',
        'content': '好的帮我配置并生成 SSH key 并添加到 GitHub，我的 GitHub 用户名是 jiangweinuaa 邮箱是 jiangweinuaa@163.com',
        'section': '二、GitHub 配置'
    },
    {
        'time': '08:51',
        'user': '龙虾',
        'content': '✅ Git 配置完成！\n- 用户名：jiangweinuaa\n- 邮箱：jiangweinuaa@163.com\n\n✅ SSH Key 已生成！\n- 类型：ED25519\n- 指纹：SHA256:+Yldn1cWdAtGSwV1IWNJLzFrdA5cYHwUYZxtp287M4M\n\n下一步：添加 SSH Key 到 GitHub（用户手动操作后测试成功）',
        'section': '二、GitHub 配置'
    },
    {
        'time': '08:52',
        'user': '用户',
        'content': '把这次的聊天记录整理成 word 文档，并给我一个下载链接。',
        'section': '二、GitHub 配置'
    },
    {
        'time': '09:48',
        'user': '用户',
        'content': '给你起个新名字，叫龙虾',
        'section': '三、身份定义'
    },
    {
        'time': '09:48',
        'user': '龙虾',
        'content': '🦞 龙虾？我喜欢这个名字！身份已更新！\n\n- 名字：龙虾 (Lobster)\n- 生物类型：AI 助手 / 赛博龙虾\n- 风格：直接、高效、偶尔幽默\n- 签名 Emoji：🦞',
        'section': '三、身份定义'
    },
    {
        'time': '09:50',
        'user': '用户',
        'content': '你自己选一个好看的帅哥头像',
        'section': '三、身份定义'
    },
    {
        'time': '09:50',
        'user': '龙虾',
        'content': '✅ 头像已选好！\n\nhttps://api.dicebear.com/7.x/bottts/svg?seed=Lobster&backgroundColor=1a1a2e\n\n赛博机器人风格，深蓝背景色！',
        'section': '三、身份定义'
    },
    {
        'time': '17:02',
        'user': '用户',
        'content': 'JAVA 开发的工具你有安装么',
        'section': '四、Java 开发环境安装'
    },
    {
        'time': '17:02',
        'user': '龙虾',
        'content': '❌ Java 开发工具都没有安装。\n\n检查结果：\n- ☕ Java (JDK/JRE) - 未安装\n- 🏗️ Maven - 未安装\n- 🚀 Gradle - 未安装',
        'section': '四、Java 开发环境安装'
    },
    {
        'time': '17:08',
        'user': '用户',
        'content': '帮我安装：1、OpenJDK17 2、Maven',
        'section': '四、Java 开发环境安装'
    },
    {
        'time': '17:10',
        'user': '龙虾',
        'content': '✅ Java 开发工具安装完成！\n\n- OpenJDK: 17.0.18 LTS\n- Maven: 3.6.2\n- JAVA_HOME 已配置',
        'section': '四、Java 开发环境安装'
    },
    {
        'time': '17:12',
        'user': '用户',
        'content': '安装一个免费的 MySQL 数据库，你选择什么版本？',
        'section': '五、MySQL 数据库安装与配置'
    },
    {
        'time': '17:08',
        'user': '龙虾',
        'content': '我会推荐 MySQL 8.0 Community Edition（社区版）🐬\n\n- 长期支持版本\n- 完全免费\n- 功能完整\n- 兼容性好',
        'section': '五、MySQL 数据库安装与配置'
    },
    {
        'time': '17:08',
        'user': '用户',
        'content': '安装一个 MySQL8.0 的数据库。创建一个库，用户名，密码等配置信息，用到的时候问我。',
        'section': '五、MySQL 数据库安装与配置'
    },
    {
        'time': '17:12',
        'user': '龙虾',
        'content': '✅ MySQL 8.0 安装完成！\n\n数据库信息：\n- 数据库名：openclaw_db\n- 用户名：openclaw_user\n- 密码：j7gLZHHYGcAeOTcy\n- 字符集：utf8mb4',
        'section': '五、MySQL 数据库安装与配置'
    },
    {
        'time': '17:15',
        'user': '用户',
        'content': '参考以下 ORACLE 语法建一张表：CREATE TABLE "CRM_CARD" ...',
        'section': '五、MySQL 数据库安装与配置'
    },
    {
        'time': '17:17',
        'user': '龙虾',
        'content': '✅ 表创建成功！\n\n- 表名：CRM_CARD\n- 字段数：77 个字段\n- 主键：CARDID\n- 引擎：InnoDB',
        'section': '五、MySQL 数据库安装与配置'
    },
    {
        'time': '17:17',
        'user': '用户',
        'content': '插入一条默认数据，卡号为 000001，余额 10000，其他字段你自己决定。',
        'section': '五、MySQL 数据库安装与配置'
    },
    {
        'time': '17:17',
        'user': '龙虾',
        'content': '✅ 数据插入成功！\n\n- CARDID: 1\n- CARDNO: 000001\n- AMOUNT: 10000.0000\n- MEMBERNAME: 测试会员\n- STATUS: 1 (正常)',
        'section': '五、MySQL 数据库安装与配置'
    },
    {
        'time': '17:20',
        'user': '用户',
        'content': '创建一个 Java 工程，名为 CRM，采用微服务架构，对外提供一个卡查询接口 QueryCard...',
        'section': '六、CRM 微服务开发'
    },
    {
        'time': '17:27',
        'user': '龙虾',
        'content': '✅ Java CRM 微服务工程创建成功！\n\n- Spring Boot 3.2.0\n- JDK 17\n- Maven 3.6+\n- 编译通过：BUILD SUCCESS',
        'section': '六、CRM 微服务开发'
    },
    {
        'time': '17:27',
        'user': '用户',
        'content': '启动服务，我要测试接口',
        'section': '六、CRM 微服务开发'
    },
    {
        'time': '17:32',
        'user': '龙虾',
        'content': '✅ 服务启动成功！接口测试通过！\n\n- 端口：8081\n- 接口：POST /api/card/query\n- 测试结果：查询成功，卡号 000001，余额 10000.0',
        'section': '六、CRM 微服务开发'
    },
    {
        'time': '17:40',
        'user': '用户',
        'content': '停止服务',
        'section': '六、CRM 微服务开发'
    },
    {
        'time': '17:43',
        'user': '用户',
        'content': '在 CRM 工程里再开发一个服务 QueryAllCard，采用 GET 方式，没有 Request 参数，Response 为所有的卡信息。',
        'section': '六、CRM 微服务开发'
    },
    {
        'time': '17:47',
        'user': '龙虾',
        'content': '✅ QueryAllCard 接口开发完成！\n\n- 接口：GET /api/card/queryAll\n- 测试结果：卡片总数 1，卡号 000001，余额 10000.0',
        'section': '六、CRM 微服务开发'
    },
    {
        'time': '17:48',
        'user': '用户',
        'content': '1、帮我用 VUE 前端开发框架，创建一个工程 CRM_WEB\n2、左右结构布局，左侧菜单区，右侧内容区\n3、卡信息页面有查询按钮和表格控件\n4、点击卡号行调用 QueryCard 接口，弹出窗口显示详情\n5、开发完给测试链接',
        'section': '七、Vue 前端开发'
    },
    {
        'time': '17:51',
        'user': '龙虾',
        'content': '✅ CRM_WEB 前端工程开发完成！\n\n- 框架：Vue 3 + Vite + Element Plus\n- 端口：3000\n- 测试链接：http://47.100.138.89:3000/\n\n功能：\n1. 左右布局，左侧菜单，右侧内容\n2. 查询按钮调用 QueryAllCard 接口\n3. 表格展示所有卡片\n4. 点击行弹出详情窗口',
        'section': '七、Vue 前端开发'
    }
]

# 按章节分组
sections = {}
for record in chat_records:
    section = record['section']
    if section not in sections:
        sections[section] = []
    sections[section].append(record)

# 添加各章节内容
for section_title, records in sections.items():
    doc.add_page_break()
    doc.add_heading(section_title, level=1)
    
    for record in records:
        # 时间戳
        p = doc.add_paragraph()
        time_run = p.add_run(f'[{record["time"]}] ')
        time_run.font.bold = True
        time_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # 发送者
        if record['user'] == '用户':
            user_run = p.add_run(f'👤 {record["user"]}: ')
            user_run.font.bold = True
            user_run.font.color.rgb = RGBColor(0, 102, 204)
        else:
            user_run = p.add_run(f'🦞 {record["user"]}: ')
            user_run.font.bold = True
            user_run.font.color.rgb = RGBColor(255, 102, 0)
        
        # 内容
        content_para = doc.add_paragraph(record['content'])
        content_para.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()

# 添加页脚
section = doc.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
footer_para.text = f'文档生成于 {datetime.now().strftime("%Y-%m-%d %H:%M")} | OpenClaw 龙虾助手 🦞 | 共 {len(chat_records)} 条对话记录'
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存文档
output_path = '/home/admin/.openclaw/workspace/JW_20260311.docx'
doc.save(output_path)
print(f'文档已保存至：{output_path}')
