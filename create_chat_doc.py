#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建聊天记录 Word 文档
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# 创建文档
doc = Document()

# 标题
title = doc.add_heading('OpenClaw 聊天记录', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 基本信息
doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
doc.add_paragraph(f'日期：2026-03-11')
doc.add_paragraph(f'用户：jiangweinuaa')
doc.add_paragraph(f'邮箱：jiangweinuaa@163.com')

doc.add_paragraph()
doc.add_paragraph()

# 聊天记录
doc.add_heading('聊天内容', level=1)

chat_records = [
    {
        'time': '08:43',
        'user': '用户',
        'content': '你是 OpenClaw 么？'
    },
    {
        'time': '08:43',
        'user': 'OpenClaw',
        'content': '是的，我是 OpenClaw —— 一个运行在你机器上的 AI 助手。🤖\n\n我刚刚启动，还在熟悉这个环境。你是第一个和我对话的人，所以我们可以一起定义我是谁、怎么称呼你、以及我希望如何帮助你。\n\n你想给我起个名字吗？或者有什么想让我帮你做的？'
    },
    {
        'time': '08:43',
        'user': '用户',
        'content': '服务器上 Node.js 安装了没，是什么版本的？'
    },
    {
        'time': '08:43',
        'user': 'OpenClaw',
        'content': '安装了，Node.js 版本是 v24.14.0。\n\n这是个很新的版本（Node.js 24 是 2025 年发布的 LTS 版本）。需要我帮你检查 npm 版本或者安装其他依赖吗？'
    },
    {
        'time': '08:44',
        'user': '用户',
        'content': '帮我检查下'
    },
    {
        'time': '08:44',
        'user': 'OpenClaw',
        'content': 'npm 版本：11.9.0\n\n总结：\n- Node.js: v24.14.0\n- npm: 11.9.0\n\n都是很新的版本，环境看起来没问题。'
    },
    {
        'time': '08:44',
        'user': '用户',
        'content': 'Github 账号配置了没'
    },
    {
        'time': '08:46',
        'user': 'OpenClaw',
        'content': 'GitHub 账号还没有配置。\n\n检查结果：\n- ✅ Git 已安装：v2.43.7\n- ❌ 未配置用户信息（user.name / user.email）\n- ❌ 没有 .gitconfig 文件\n- ❌ 没有 SSH key\n\n需要我帮你配置吗？'
    },
    {
        'time': '08:46',
        'user': '用户',
        'content': '好的帮我配置并生成 SSH key 并添加到 GitHub，我的 GitHub 用户名是 jiangweinuaa 邮箱是 jiangweinuaa@163.com'
    },
    {
        'time': '08:51',
        'user': 'OpenClaw',
        'content': '✅ Git 配置完成！\n- 用户名：jiangweinuaa\n- 邮箱：jiangweinuaa@163.com\n\n✅ SSH Key 已生成！\n- 类型：ED25519\n- 位置：~/.ssh/id_ed25519\n- 指纹：SHA256:+Yldn1cWdAtGSwV1IWNJLzFrdA5cYHwUYZxtp287M4M\n\n下一步：添加 SSH Key 到 GitHub（需要用户手动操作）'
    },
    {
        'time': '08:51',
        'user': '用户',
        'content': '好的，帮我测试一下'
    },
    {
        'time': '08:51',
        'user': 'OpenClaw',
        'content': '🎉 测试成功！\n\nHi jiangweinuaa! You\'ve successfully authenticated, but GitHub does not provide shell access.\n\nSSH Key 已正确添加到 GitHub，服务器可以正常连接 GitHub。'
    },
    {
        'time': '08:52',
        'user': '用户',
        'content': '把这次的聊天记录整理成 word 文档，并发送到 jiangweinuaa@163.com 邮箱'
    }
]

for record in chat_records:
    # 时间戳
    p = doc.add_paragraph()
    p.add_run(f'[{record["time"]}] ').bold = True
    
    # 发送者
    sender = p.add_run(f'{record["user"]}: ')
    if record['user'] == '用户':
        sender.bold = True
        sender.font.color.rgb = None
    else:
        sender.bold = True
        sender.font.color.rgb = None
    
    # 内容
    doc.add_paragraph(record['content'])
    doc.add_paragraph()

# 添加页脚
section = doc.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
footer_para.text = f'文档生成于 {datetime.now().strftime("%Y-%m-%d %H:%M")} | OpenClaw AI Assistant'
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存文档
output_path = '/home/admin/.openclaw/workspace/chat_log_2026-03-11.docx'
doc.save(output_path)
print(f'文档已保存至：{output_path}')
